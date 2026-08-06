"""
Core conversion logic for the Excel -> Word report builder.

Design:
- Read an uploaded .xlsx and detect its column headers.
- The user supplies `column_order`: a list of column names in the order
  they want them to appear.
- For every DATA ROW in the sheet, generate one repeated "block" in the
  Word document. Within each block, each chosen column becomes its own
  paragraph, in the user's chosen order.
- The FIRST column in the user's order is treated as the block's heading
  (e.g. a person's name), the rest are rendered as "Label: value" lines.
  This is a sensible default and easy to change later.
"""

from docx import Document
from docx.shared import Pt
import pandas as pd


def get_sheets(excel_path: str) -> list[str]:
    """List all sheet names in the workbook, so the user can pick one."""
    xl = pd.ExcelFile(excel_path)
    return xl.sheet_names


def get_columns(excel_path: str, sheet_name) -> list[str]:
    """Detect column headers from a specific sheet."""
    df = pd.read_excel(excel_path, sheet_name=sheet_name, nrows=0)
    return list(df.columns)


def build_report(excel_path: str, sheet_name, heading_column: str,
                  column_order: list[str], output_path: str,
                  title: str | None = None) -> str:
    """
    Build a Word report from one sheet of an Excel file.

    excel_path:      path to the uploaded .xlsx
    sheet_name:       which sheet to read (from get_sheets)
    heading_column:   column the user picked as the heading for each block
                       (kept separate from column_order — may or may not
                       also appear in column_order, see note below)
    column_order:     ordered list of column names for the body of each
                       block, in the order the user arranged them
    output_path:      where to write the resulting .docx
    title:            optional document title
    """
    df = pd.read_excel(excel_path, sheet_name=sheet_name)

    if heading_column not in df.columns:
        raise ValueError(f"Heading column not found in sheet: {heading_column}")
    missing = [c for c in column_order if c not in df.columns]
    if missing:
        raise ValueError(f"Columns not found in sheet: {missing}")

    # Body columns = the user's chosen order, minus the heading column
    # itself (no point repeating it as a labeled line under its own heading)
    body_cols = [c for c in column_order if c != heading_column]

    doc = Document()

    if title:
        doc.add_heading(title, level=0)

    for i, row in df.iterrows():
        # Heading for this row's block
        heading_value = "" if pd.isna(row[heading_column]) else str(row[heading_column])
        doc.add_heading(heading_value, level=1)

        # Remaining chosen columns as labeled body paragraphs, in order
        for col in body_cols:
            value = "" if pd.isna(row[col]) else str(row[col])
            p = doc.add_paragraph()
            label_run = p.add_run(f"{col}: ")
            label_run.bold = True
            p.add_run(value)

        # Spacer between entries (skip after the last row)
        if i < len(df) - 1:
            doc.add_paragraph()

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    sheets = get_sheets("sample_input.xlsx")
    print("Sheets:", sheets)

    cols = get_columns("sample_input.xlsx", sheets[0])
    print("Detected columns:", cols)

    # Simulate user choices: heading = Name, body order = Role, Region, Bio
    # (Email deliberately left out, to prove selection/ordering both work)
    heading_column = "Name"
    user_order = ["Role", "Region", "Bio"]

    build_report(
        "sample_input.xlsx",
        sheets[0],
        heading_column,
        user_order,
        "sample_output.docx",
        title="Team Profiles Report",
    )
    print("Wrote sample_output.docx")
