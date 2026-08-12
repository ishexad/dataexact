"""
FastAPI router + branding engine for the Format Report tool.

Flow (single request, no intermediate steps needed):
1. POST /format/generate -> given a plain .docx, a theme id, and optionally
   a logo image + a cover title, returns a branded .docx: cover page,
   themed heading/body fonts and colors, a branded header/footer, and a
   page border.

Design: branding is done by restyling the document's own style objects
(Normal, Heading 1/2/3, Title) rather than touching individual paragraphs,
so any paragraph in the uploaded document that already uses those styles
picks up the new look automatically. Paragraphs with direct/manual
formatting (e.g. someone bolded text by hand instead of using a Heading
style) are left as-is — this tool assumes a "plain" document uses Word's
normal semantic styles, which is the common case.
"""

import io
import os
import uuid

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse

import gating

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"
MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB, same cap as the other tools
MAX_LOGO_BYTES = 2 * 1024 * 1024

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

router = APIRouter(prefix="/format")

THEMES = {
    "corporate": {
        "name": "Corporate",
        "description": "Navy, traditional, built for reports.",
        "accent": "1F4E79",
        "heading_font": "Cambria",
        "body_font": "Calibri",
    },
    "minimal": {
        "name": "Minimal",
        "description": "Charcoal, restrained, lots of whitespace.",
        "accent": "3B3B3B",
        "heading_font": "Calibri Light",
        "body_font": "Calibri",
    },
    "bold": {
        "name": "Bold",
        "description": "Warm accent, confident headings.",
        "accent": "C2661C",
        "heading_font": "Arial",
        "body_font": "Arial",
    },
}


def _add_border(pPr, edge: str, color_hex: str, size: str = "12", space: str = "4"):
    """Add a single-edge border to a paragraph's pPr (used for header/footer
    rules and the cover page underline). python-docx has no high-level API
    for paragraph borders, so this builds the raw <w:pBdr> XML directly."""
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    border = OxmlElement(f"w:{edge}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color_hex)
    pBdr.append(border)


def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _clear_paragraph_runs(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _apply_typography(doc: Document, theme: dict):
    styles = doc.styles
    accent = RGBColor.from_string(theme["accent"])

    try:
        normal = styles["Normal"]
        normal.font.name = theme["body_font"]
        normal.font.size = Pt(11)
    except KeyError:
        pass

    for level, size in ((1, 20), (2, 16), (3, 13)):
        try:
            style = styles[f"Heading {level}"]
        except KeyError:
            continue
        style.font.name = theme["heading_font"]
        style.font.size = Pt(size)
        style.font.color.rgb = accent
        style.font.bold = True

    try:
        title_style = styles["Title"]
        title_style.font.name = theme["heading_font"]
        title_style.font.size = Pt(28)
        title_style.font.color.rgb = accent
        title_style.font.bold = True
    except KeyError:
        pass


def _apply_page_layout(doc: Document, theme: dict):
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    sect_pr = section._sectPr
    for existing in sect_pr.findall(qn("w:pgBorders")):
        sect_pr.remove(existing)

    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "18")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), theme["accent"])
        pg_borders.append(border)
    sect_pr.append(pg_borders)


def _apply_header_footer(doc: Document, theme: dict, logo_bytes: bytes | None):
    section = doc.sections[0]
    accent = theme["accent"]

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    _clear_paragraph_runs(hp)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_bytes:
        hp.add_run().add_picture(io.BytesIO(logo_bytes), height=Inches(0.4))
    else:
        run = hp.add_run(theme["name"])
        run.font.name = theme["heading_font"]
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(accent)
    _add_border(hp._p.get_or_add_pPr(), "bottom", accent)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    _clear_paragraph_runs(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(fp)
    for run in fp.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(accent)


def _apply_cover_page(doc: Document, theme: dict, title: str, logo_bytes: bytes | None):
    body = doc.element.body
    accent = RGBColor.from_string(theme["accent"])
    cover_paragraphs = []

    if logo_bytes:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(io.BytesIO(logo_bytes), height=Inches(1.2))
        cover_paragraphs.append(p_logo)

    cover_paragraphs.append(doc.add_paragraph())  # spacer

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.font.name = theme["heading_font"]
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = accent
    cover_paragraphs.append(p_title)

    p_rule = doc.add_paragraph()
    _add_border(p_rule._p.get_or_add_pPr(), "bottom", theme["accent"], size="18", space="1")
    cover_paragraphs.append(p_rule)

    p_break = doc.add_paragraph()
    p_break.add_run().add_break(WD_BREAK.PAGE)
    cover_paragraphs.append(p_break)

    # These paragraphs were just appended after the document's existing
    # content — move them to the front of the body, in the same relative
    # order, so the cover page renders before everything else.
    elements = [p._p for p in cover_paragraphs]
    for element in elements:
        body.remove(element)
    for element in reversed(elements):
        body.insert(0, element)


def _fallback_title(filename: str) -> str:
    base = os.path.splitext(os.path.basename(filename or ""))[0]
    cleaned = base.replace("_", " ").replace("-", " ").strip()
    return cleaned.title() if cleaned else "Report"


def apply_branding(input_path: str, theme_id: str, output_path: str,
                    logo_bytes: bytes | None = None, title: str | None = None,
                    filename_hint: str | None = None) -> str:
    if theme_id not in THEMES:
        raise ValueError(f"Unknown theme: {theme_id}")
    theme = THEMES[theme_id]

    doc = Document(input_path)
    cover_title = title.strip() if title and title.strip() else _fallback_title(filename_hint)

    _apply_typography(doc, theme)
    _apply_page_layout(doc, theme)
    _apply_header_footer(doc, theme, logo_bytes)
    _apply_cover_page(doc, theme, cover_title, logo_bytes)

    doc.save(output_path)
    return output_path


@router.get("/themes")
async def list_themes():
    return [{"id": key, **theme} for key, theme in THEMES.items()]


@router.post("/generate")
async def generate(
    request: Request,
    file: UploadFile = File(...),
    theme: str = Form(...),
    title: str = Form(None),
    logo: UploadFile = File(None),
):
    if theme not in THEMES:
        raise HTTPException(400, f"Unknown theme: {theme}")

    logged_in = gating.gate(request, "format")

    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "Please upload a .docx file")
    contents = await file.read()
    if len(contents) > MAX_UPLOAD_BYTES:
        raise HTTPException(400, "File too large (10 MB limit)")

    logo_bytes = None
    if logo is not None and logo.filename:
        logo_bytes = await logo.read()
        if len(logo_bytes) > MAX_LOGO_BYTES:
            raise HTTPException(400, "Logo image too large (2 MB limit)")

    file_id = str(uuid.uuid4())
    input_path = os.path.join(UPLOAD_DIR, f"{file_id}.docx")
    with open(input_path, "wb") as f:
        f.write(contents)

    output_id = str(uuid.uuid4())
    output_path = os.path.join(OUTPUT_DIR, f"{output_id}.docx")

    try:
        apply_branding(input_path, theme, output_path, logo_bytes=logo_bytes,
                        title=title, filename_hint=file.filename)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception:
        raise HTTPException(400, "Could not read this file as a valid Word document")

    gating.record_use(request, "format", logged_in)

    return FileResponse(
        output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="branded_report.docx",
    )
