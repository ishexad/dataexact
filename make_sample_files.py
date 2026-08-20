"""Build the try-it-with-a-sample files served at /samples/.

Two already existed in the repo root as developer fixtures; the other three
are written here. Each is built to actually demonstrate its own tool — the
anomaly sheet in particular carries one instance of every issue the detectors
look for, so a first-time visitor sees a full result rather than an empty one.

    python make_sample_files.py

Writes static/marketing/samples/, which the marketing mount serves from
/samples/. Rerun it if the detectors change — the member list is tuned to trip
every one of them, and a sample that no longer demonstrates the tool is worse
than no sample at all. Worth re-checking the output before committing:

    python -c "import anomaly_detector as a; print(a.detect('static/marketing/samples/member-list.xlsx','Members')['summary'])"

Needs openpyxl and python-docx, both already app dependencies.
"""
import os
import shutil

from docx import Document
from openpyxl import Workbook

# Run from anywhere: paths below are relative to the repo root.
os.chdir(os.path.dirname(os.path.abspath(__file__)))
OUT = "static/marketing/samples"
os.makedirs(OUT, exist_ok=True)

# --- Excel to Word / Compare: reuse the existing fixtures -----------------
shutil.copy("sample_input.xlsx", f"{OUT}/team-directory.xlsx")
shutil.copy("sample_a.xlsx", f"{OUT}/price-list-january.xlsx")
shutil.copy("sample_b.xlsx", f"{OUT}/price-list-february.xlsx")

# --- Detect Anomalies: one of everything the detectors find ---------------
wb = Workbook()
ws = wb.active
ws.title = "Members"
ws.append(["Member ID", "Full Name", "Email", "Joined", "Annual Fee", "Region"])

rows = [
    ["M-1001", "Alice Okafor", "alice@example.com", "2024-03-11", 240, "North"],
    ["M-1002", "Ben Carter", "ben@example.com", "2024-03-14", 265, "North"],
    ["M-1003", "Priya Nair", "priya@example.com", "2024-04-02", 195, "South"],
    # duplicated identifier, in a column that should be unique
    ["M-1003", "Priya Nair ", "priya.n@example.com", "2024-04-02", 240, "South"],
    # trailing whitespace + inconsistent capitalisation
    ["M-1004", "DAVID OYELARAN", "david@example.com", "2024-04-19", 310, "  East"],
    # a date written the American way in a column that is otherwise ISO
    ["M-1005", "Sofia Marino", "sofia@example.com", "05/22/2024", 240, "East"],
    # an outlier fee
    ["M-1006", "Tom Whitfield", "tom@example.com", "2024-05-30", 9600, "West"],
    # missing value in a column that is otherwise always filled
    ["M-1007", "Grace Liu", "", "2024-06-08", 275, "West"],
    # text where the column is otherwise numeric
    ["M-1008", "Hassan Ali", "hassan@example.com", "2024-06-21", "TBC", "North"],
    ["M-1009", "Mia Fenwick", "mia@example.com", "2024-07-03", 220, "South"],
    ["M-1010", "Noah Brandt", "noah@example.com", "2024-07-17", 290, "South"],
    # an exact duplicate of the row above — must match it in every column
    ["M-1010", "Noah Brandt", "noah@example.com", "2024-07-17", 290, "South"],
    ["M-1011", "Ayesha Khan", "ayesha@example.com", "2024-08-01", 205, "East"],
    # double space inside a value
    ["M-1012", "Liam  Doyle", "liam@example.com", "2024-08-15", 255, "West"],
]
for r in rows:
    ws.append(r)
for col, width in zip("ABCDEF", (12, 20, 26, 14, 12, 10)):
    ws.column_dimensions[col].width = width
wb.save(f"{OUT}/member-list.xlsx")

# --- Format Report: a plain document using real Word styles ---------------
doc = Document()
doc.add_heading("Quarterly Operations Review", level=0)
doc.add_paragraph(
    "This document is a sample. It uses Word's built-in heading styles and "
    "nothing else, which is exactly the kind of plain document the Format "
    "Report tool is built to brand."
)

doc.add_heading("Summary", level=1)
doc.add_paragraph(
    "Delivery held steady through the quarter. Volumes rose in the North and "
    "East regions while the South was flat, and no service level was missed."
)

doc.add_heading("Performance", level=1)
doc.add_heading("Volumes", level=2)
doc.add_paragraph(
    "Total processed volume was 12,480 units against a forecast of 11,900. "
    "The variance sat almost entirely in the North region."
)
doc.add_heading("Turnaround", level=2)
doc.add_paragraph(
    "Median turnaround was 2.4 days, unchanged on the previous quarter. The "
    "slowest decile improved from 9 days to 7."
)

doc.add_heading("Risks", level=1)
for risk in (
    "Supplier concentration in the East remains above the agreed threshold.",
    "Two of the four regional leads are on fixed-term contracts expiring in March.",
    "The manual reconciliation step is a single point of failure.",
):
    doc.add_paragraph(risk, style="List Bullet")

doc.add_heading("Next quarter", level=1)
doc.add_paragraph(
    "Priorities are to onboard a second supplier in the East, convert the two "
    "fixed-term leads to permanent contracts, and automate the reconciliation step."
)
doc.save(f"{OUT}/quarterly-review.docx")

# --- Qualitative Coding: a small corpus with recurring themes -------------
INTERVIEWS = {
    "interview-site-a.docx": (
        "Interview — Site A",
        [
            "The cost of the new system was the first thing people raised. Several "
            "staff said the budget had been agreed before anyone asked what the work "
            "actually involved, and the price kept moving.",
            "Training was thin. Two afternoons, and then you were expected to know it. "
            "The people who had used something similar before coped; everyone else "
            "learned from colleagues.",
            "On balance the team felt the change was worth it, but they were clear "
            "that the timeline had been unrealistic and the workload during the "
            "transition was not acknowledged.",
        ],
    ),
    "interview-site-b.docx": (
        "Interview — Site B",
        [
            "Staff here were more positive about training. There was a proper "
            "programme, with follow-up sessions, and people said that made the "
            "difference to how confident they felt.",
            "Cost came up again, though framed differently: not the price itself but "
            "the lack of transparency about what the budget covered.",
            "Workload during the transition was the dominant complaint. Several "
            "people described working evenings for a month with no recognition.",
        ],
    ),
    "interview-site-c.docx": (
        "Interview — Site C",
        [
            "This site had the smoothest experience. Training was run by someone from "
            "their own team, which staff repeatedly said mattered more than the "
            "content of the training itself.",
            "There was little discussion of cost. When asked directly, most people "
            "said they had no visibility of the budget and did not consider it their "
            "concern.",
            "Workload was manageable because the timeline was extended twice. Staff "
            "contrasted this with what they had heard from other sites.",
        ],
    ),
}
for filename, (title, paragraphs) in INTERVIEWS.items():
    d = Document()
    d.add_heading(title, level=1)
    for p in paragraphs:
        d.add_paragraph(p)
    d.save(f"{OUT}/{filename}")

for name in sorted(os.listdir(OUT)):
    print("  %-32s %6.1f KB" % (name, os.path.getsize(f"{OUT}/{name}") / 1024))
