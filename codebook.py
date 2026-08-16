"""
FastAPI router for the Qualitative Coding tool (Dedoose-style, rule-based).

Flow:
1. POST /codebook/upload   -> upload one or more articles (.pdf / .docx),
                               get back a doc_id per file
2. POST /codebook/extract  -> given those doc_ids (each with an optional
                               tag) + a code tree, get back every coded
                               excerpt across every file, plus per-code
                               frequency counts
3. POST /codebook/export   -> same inputs, returned as an .xlsx workbook

The coding model mirrors Dedoose: a two-level code tree (main codes, each
with optional sub-codes) is applied across a whole corpus, and each match
becomes an *excerpt* — the passage itself, plus which file/tag/location it
came from. One code can produce many excerpts per document; that's the
point. Analysis then runs off excerpt counts per code per document.

Matching is deterministic keyword search — no LLM, no API key. Each code
carries the keywords that signal it; a code with no keywords falls back to
matching its own name. Single-word keywords match on word boundaries so
"cost" doesn't fire on "costume"; multi-word keywords match as phrases.

Documents are segmented into passage-sized chunks first (PDF text blocks,
Word paragraphs/table cells) so an excerpt is a readable unit rather than
a whole page or a fragment of a sentence.
"""

import io
import json
import os
import re
import uuid
from typing import Dict, List, Optional

import pymupdf
from docx import Document
from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from pydantic import BaseModel

import analytics
import gating

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

router = APIRouter(prefix="/codebook")

MAX_FILES = 25              # per upload request
MAX_SEGMENT_CHARS = 1200    # longer passages get split into sentence groups
MAX_EXCERPTS = 5000         # hard cap so a broad code can't produce a runaway response

# How long an uploaded corpus survives before main.py's cleanup sweep removes
# it. The other tools are one-file-in, one-file-out and finish in a minute, so
# an hour suits them; coding is different — tagging 25 files and building a
# code tree is real work that can easily outlast an hour, and losing the corpus
# mid-session would mean re-uploading everything. Affordable because the
# uploaded PDFs/Word files are never kept: only their extracted text is, as
# JSON, which is a fraction of the size. main.py reads this for *.json in
# uploads/, and /upload reports it so the UI can't quote a stale number.
CORPUS_MAX_AGE_SECONDS = 6 * 60 * 60

HEADER_FILL = PatternFill(start_color="2E2E2E", end_color="2E2E2E", fill_type="solid")
HEADER_FONT = Font(color="FFFFFF", bold=True)

SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?])\s+")
SINGLE_TOKEN_RE = re.compile(r"[\w\-']+")


class SubCode(BaseModel):
    name: str
    keywords: List[str] = []


class Code(BaseModel):
    name: str
    keywords: List[str] = []
    subcodes: List[SubCode] = []


class CodeTree(BaseModel):
    codes: List[Code]


class DocumentRef(BaseModel):
    doc_id: str
    tag: str = ""


class DocumentRefs(BaseModel):
    documents: List[DocumentRef]


def _doc_path(doc_id: str) -> str:
    return os.path.join(UPLOAD_DIR, f"{doc_id}.json")


# ---------------------------------------------------------------------
# Document parsing -> passage-sized segments
# ---------------------------------------------------------------------

def _clean(text: str) -> str:
    """Collapse the hard line wrapping that PDF/Word extraction leaves behind
    so an excerpt reads as prose rather than as broken lines."""
    return re.sub(r"\s+", " ", text or "").strip()


def _split_long(text: str) -> List[str]:
    """Keep excerpts readable: anything over MAX_SEGMENT_CHARS is regrouped
    into sentence-aligned chunks rather than cut mid-word."""
    if len(text) <= MAX_SEGMENT_CHARS:
        return [text]

    chunks, current = [], ""
    for sentence in SENTENCE_SPLIT_RE.split(text):
        if current and len(current) + len(sentence) + 1 > MAX_SEGMENT_CHARS:
            chunks.append(current.strip())
            current = sentence
        else:
            current = f"{current} {sentence}".strip()
    if current.strip():
        chunks.append(current.strip())
    return chunks


def _segments_from_pdf(file_bytes: bytes) -> List[dict]:
    """Use PyMuPDF's text *blocks* rather than raw page text — blocks are
    already paragraph-shaped, which is what we want an excerpt to be."""
    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    segments = []
    for page_num, page in enumerate(doc, start=1):
        for block in page.get_text("blocks"):
            if len(block) >= 7 and block[6] != 0:
                continue  # image block, not text
            text = _clean(block[4])
            if not text:
                continue
            for piece in _split_long(text):
                segments.append({"page": page_num, "location": f"p. {page_num}", "text": piece})
    doc.close()
    return segments


def _segments_from_docx(file_bytes: bytes) -> List[dict]:
    """Word has no page numbers available without rendering, so cite the
    paragraph/table position instead."""
    document = Document(io.BytesIO(file_bytes))
    segments = []

    for para_num, para in enumerate(document.paragraphs, start=1):
        text = _clean(para.text)
        if not text:
            continue
        for piece in _split_long(text):
            segments.append({"page": None, "location": f"para. {para_num}", "text": piece})

    for table_num, table in enumerate(document.tables, start=1):
        for row_num, row in enumerate(table.rows, start=1):
            for col_num, cell in enumerate(row.cells, start=1):
                text = _clean(cell.text)
                if not text:
                    continue
                location = f"table {table_num}, r{row_num}c{col_num}"
                for piece in _split_long(text):
                    segments.append({"page": None, "location": location, "text": piece})

    return segments


@router.post("/upload")
async def upload(request: Request, files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(400, "Please upload at least one file")
    if len(files) > MAX_FILES:
        raise HTTPException(400, f"Too many files at once ({MAX_FILES} maximum)")

    limit = gating.max_upload_bytes(request)
    documents = []
    # A corpus arrives as one request of many files, so one upload event
    # covers the batch, sized by their combined bytes.
    total_bytes = 0

    def _fail(message: str):
        analytics.log_event(request, "codebook", analytics.EVENT_UPLOAD,
                            file_size_bytes=total_bytes, success=False,
                            error_message=message)
        return HTTPException(400, message)

    for file in files:
        name_lower = (file.filename or "").lower()
        if not name_lower.endswith((".pdf", ".docx")):
            raise _fail(f"'{file.filename}' isn't a .pdf or .docx file")

        contents = await file.read()
        total_bytes += len(contents)
        if len(contents) > limit:
            raise _fail(f"'{file.filename}' is too large ({limit // (1024 * 1024)} MB limit)")

        try:
            if name_lower.endswith(".pdf"):
                segments = _segments_from_pdf(contents)
            else:
                segments = _segments_from_docx(contents)
        except Exception:
            raise _fail(f"Could not read '{file.filename}' — it may be corrupt or password-protected")

        if not segments:
            raise _fail(
                f"No extractable text found in '{file.filename}' (a scanned PDF needs OCR first)"
            )

        doc_id = str(uuid.uuid4())
        with open(_doc_path(doc_id), "w", encoding="utf-8") as f:
            json.dump({"filename": file.filename, "segments": segments}, f)

        documents.append({
            "doc_id": doc_id,
            "filename": file.filename,
            "segment_count": len(segments),
        })

    analytics.log_event(request, "codebook", analytics.EVENT_UPLOAD, file_size_bytes=total_bytes)
    return {
        "documents": documents,
        "expires_in_hours": CORPUS_MAX_AGE_SECONDS // 3600,
    }


# ---------------------------------------------------------------------
# Coding — apply the code tree across every document
# ---------------------------------------------------------------------

def _compile_keyword(keyword: str) -> Optional[re.Pattern]:
    keyword = keyword.strip()
    if not keyword:
        return None
    # Single words match on word boundaries ("cost" shouldn't hit "costume");
    # phrases match literally, since boundaries around punctuation get messy.
    if SINGLE_TOKEN_RE.fullmatch(keyword):
        return re.compile(rf"\b{re.escape(keyword)}\b", re.IGNORECASE)
    return re.compile(re.escape(keyword), re.IGNORECASE)


def _flatten_codes(tree: CodeTree) -> List[dict]:
    """Turn the two-level tree into a flat list of matchable entries. A main
    code is matchable in its own right (Dedoose lets you apply a parent code
    directly), and each sub-code is matchable separately."""
    entries = []
    for code in tree.codes:
        code_name = code.name.strip()
        if not code_name:
            continue

        keywords = [k for k in code.keywords if k.strip()] or [code_name]
        patterns = [(k, _compile_keyword(k)) for k in keywords]
        entries.append({
            "code": code_name,
            "subcode": "",
            "patterns": [(k, p) for k, p in patterns if p],
        })

        for sub in code.subcodes:
            sub_name = sub.name.strip()
            if not sub_name:
                continue
            sub_keywords = [k for k in sub.keywords if k.strip()] or [sub_name]
            sub_patterns = [(k, _compile_keyword(k)) for k in sub_keywords]
            entries.append({
                "code": code_name,
                "subcode": sub_name,
                "patterns": [(k, p) for k, p in sub_patterns if p],
            })

    return entries


def _load_documents(refs: List[DocumentRef]) -> List[dict]:
    documents = []
    for ref in refs:
        path = _doc_path(ref.doc_id)
        if not os.path.exists(path):
            raise HTTPException(404, "One of the files has expired — please re-upload")
        with open(path, "r", encoding="utf-8") as f:
            stored = json.load(f)
        documents.append({
            "doc_id": ref.doc_id,
            "filename": stored["filename"],
            "tag": (ref.tag or "").strip(),
            "segments": stored["segments"],
        })
    return documents


def _run_coding(documents: List[dict], entries: List[dict]) -> dict:
    excerpts = []
    # counts[(code, subcode)][doc_id] -> number of excerpts
    counts: Dict[tuple, Dict[str, int]] = {
        (e["code"], e["subcode"]): {} for e in entries
    }
    truncated = False

    for document in documents:
        for segment in document["segments"]:
            text = segment["text"]
            for entry in entries:
                matched = [kw for kw, pattern in entry["patterns"] if pattern.search(text)]
                if not matched:
                    continue

                if len(excerpts) >= MAX_EXCERPTS:
                    truncated = True
                    break

                excerpts.append({
                    "doc_id": document["doc_id"],
                    "filename": document["filename"],
                    "tag": document["tag"],
                    "code": entry["code"],
                    "subcode": entry["subcode"],
                    "page": segment["page"],
                    "location": segment["location"],
                    "text": text,
                    "keywords": matched,
                })
                key = (entry["code"], entry["subcode"])
                counts[key][document["doc_id"]] = counts[key].get(document["doc_id"], 0) + 1
            if truncated:
                break
        if truncated:
            break

    code_counts = [
        {
            "code": code,
            "subcode": subcode,
            "total": sum(by_doc.values()),
            "by_document": by_doc,
        }
        for (code, subcode), by_doc in counts.items()
    ]

    documents_out = [
        {
            "doc_id": d["doc_id"],
            "filename": d["filename"],
            "tag": d["tag"],
            "excerpt_count": sum(1 for e in excerpts if e["doc_id"] == d["doc_id"]),
        }
        for d in documents
    ]

    return {
        "documents": documents_out,
        "code_counts": code_counts,
        "excerpts": excerpts,
        "summary": {
            "documents": len(documents_out),
            "codes": len(code_counts),
            "excerpts": len(excerpts),
            "truncated": truncated,
        },
    }


def _parse_inputs(documents_json: str, codes_json: str) -> tuple:
    try:
        refs = DocumentRefs(**json.loads(documents_json))
    except Exception as e:
        raise HTTPException(400, f"Invalid documents: {e}")
    if not refs.documents:
        raise HTTPException(400, "Upload at least one file first")

    try:
        tree = CodeTree(**json.loads(codes_json))
    except Exception as e:
        raise HTTPException(400, f"Invalid codes: {e}")

    entries = _flatten_codes(tree)
    if not entries:
        raise HTTPException(400, "Add at least one code with a name")

    return refs.documents, entries


@router.post("/extract")
async def extract(
    request: Request,
    documents_json: str = Form(...),
    codes_json: str = Form(...),
):
    logged_in = gating.gate(request, "codebook")
    try:
        refs, entries = _parse_inputs(documents_json, codes_json)
        documents = _load_documents(refs)
        results = _run_coding(documents, entries)
    except HTTPException as e:
        analytics.log_event(request, "codebook", analytics.EVENT_PROCESSED,
                            success=False, error_message=str(e.detail))
        raise
    except Exception as e:
        analytics.log_event(request, "codebook", analytics.EVENT_PROCESSED,
                            success=False, error_message=str(e))
        raise

    analytics.log_event(request, "codebook", analytics.EVENT_PROCESSED)
    gating.record_use(request, "codebook", logged_in)
    return results


# ---------------------------------------------------------------------
# Excel export
# ---------------------------------------------------------------------

def _autosize(ws):
    for col_cells in ws.columns:
        length = max((len(str(c.value)) if c.value is not None else 0) for c in col_cells)
        ws.column_dimensions[get_column_letter(col_cells[0].column)].width = min(max(length + 2, 10), 60)


def _write_header(ws, headers, row=1):
    for i, header in enumerate(headers, start=1):
        cell = ws.cell(row=row, column=i, value=header)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(vertical="center")


def build_export(results: dict, output_path: str) -> str:
    wb = Workbook()

    # --- Summary ---
    ws = wb.active
    ws.title = "Summary"
    ws.append(["Qualitative Coding Summary"])
    ws["A1"].font = Font(bold=True, size=14)
    ws.append([])
    ws.append(["Documents coded", results["summary"]["documents"]])
    ws.append(["Codes applied", results["summary"]["codes"]])
    ws.append(["Excerpts found", results["summary"]["excerpts"]])
    if results["summary"]["truncated"]:
        ws.append(["Note", f"Results capped at {MAX_EXCERPTS} excerpts — narrow your keywords"])
    ws.append([])
    _write_header(ws, ["Document", "Tag", "Excerpts"], row=ws.max_row + 1)
    for doc in results["documents"]:
        ws.append([doc["filename"], doc["tag"], doc["excerpt_count"]])
    _autosize(ws)

    # --- Excerpts: the coded data itself, one row per excerpt ---
    ws = wb.create_sheet("Excerpts")
    _write_header(ws, ["Document", "Tag", "Code", "Sub-code", "Location", "Matched Keywords", "Excerpt"])
    for i, excerpt in enumerate(results["excerpts"], start=2):
        ws.cell(row=i, column=1, value=excerpt["filename"])
        ws.cell(row=i, column=2, value=excerpt["tag"])
        ws.cell(row=i, column=3, value=excerpt["code"])
        ws.cell(row=i, column=4, value=excerpt["subcode"])
        ws.cell(row=i, column=5, value=excerpt["location"])
        ws.cell(row=i, column=6, value=", ".join(excerpt["keywords"]))
        ws.cell(row=i, column=7, value=excerpt["text"]).alignment = Alignment(wrap_text=True, vertical="top")
    _autosize(ws)
    ws.column_dimensions["G"].width = 80

    # --- Code Frequency: codes down the side, documents across the top ---
    ws = wb.create_sheet("Code Frequency")
    doc_headers = [d["filename"] for d in results["documents"]]
    _write_header(ws, ["Code", "Sub-code", "Total"] + doc_headers)
    for i, row in enumerate(results["code_counts"], start=2):
        ws.cell(row=i, column=1, value=row["code"])
        ws.cell(row=i, column=2, value=row["subcode"])
        ws.cell(row=i, column=3, value=row["total"])
        for j, doc in enumerate(results["documents"], start=4):
            ws.cell(row=i, column=j, value=row["by_document"].get(doc["doc_id"], 0))
    _autosize(ws)

    wb.save(output_path)
    return output_path


@router.post("/export")
async def export(
    request: Request,
    documents_json: str = Form(...),
    codes_json: str = Form(...),
):
    # Gate but don't consume: the extraction that produced this data already
    # counted against the user's allowance, so downloading it shouldn't
    # charge a second time — while still blocking anyone out of quota who
    # tries to skip /extract and call /export directly.
    gating.gate(request, "codebook")
    refs, entries = _parse_inputs(documents_json, codes_json)
    documents = _load_documents(refs)

    results = _run_coding(documents, entries)

    output_path = os.path.join(OUTPUT_DIR, f"{uuid.uuid4()}.xlsx")
    build_export(results, output_path)

    analytics.log_event(request, "codebook", analytics.EVENT_DOWNLOAD)
    return FileResponse(
        output_path,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename="coded_excerpts.xlsx",
    )
